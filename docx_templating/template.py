#!/bin/python3

from __future__ import unicode_literals, absolute_import, print_function

import sys, time, urllib, os.path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt

class Report(object):
    """Report that needs to be generated.

    Attributes:
        title: Document title.
        subtitle: Document sub-title.
        version: Version of the document. Default it '0.1'.
    """
    
    def __init__(self, title='', subtitle='', version='0.1'):
        """
        Returns a Report object whose document_name is *document_name*.
        """

        self.document = Document("template/template.docx")
        
        ######## Front page, editing the existing template
        for paragraph in self.document.paragraphs:
            if '{title}' in paragraph.text:
                paragraph.text = title
            elif '{sub-title}' in paragraph.text:
                paragraph.text = subtitle
            elif '{version}' in paragraph.text:
                paragraph.text = 'Version ' + version
            elif '{date}' in paragraph.text:
                paragraph.text = time.strftime("%d %B %Y")
                
        self.document.add_page_break()
            
            
    def add_page_break(self):
        """
        Adds a page break to the report.
        """
        self.document.add_page_break()
        
        
    def add_paragraph(self, text, bold=False, underline=False):
        """
        Adds a paragraph to the report.
        
        Parameters
        ----------
        text : str
            Text to be contained within the paragraph.
        bold : bool
            Bolds the paragraph.
        underline : bool
            Underlines the paragraph.
        """
        
        paragraph = self.document.add_paragraph().add_run(text)
        paragraph.bold = bold
        paragraph.underline = underline
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
    def delete_paragraph(self, text):
        """
        Deletes a paragraph that contains the provided text.
        
        Parameters
        ----------
        text : str
            Text to be contained within the paragraph.
        """
        
        for index, paragraph in enumerate(self.document.paragraphs):
            if text in paragraph.text:
                delete_element(paragraph)
            
            
    def add_heading(self, text, level=1, numbering=True):
        """
        Adds a heading to the report.
        
        Parameters
        ----------
        text : str
            Heading title.
        level : int
            Supported levels: 1-9.
        numbering : bool
            Show numbering in on the heading.
        """
        
        if numbering == True:
            self.document.add_paragraph(text, style="Heading " + str(level))
        else:
            self.document.add_paragraph(text, style="Heading " + str(level) + "-No Numbers")
            
    
    def add_command(self, text):
        """
        Adds a command (or command block) to the report.
        
        Parameters
        ----------
        text : str
            Command string. Separate the commands with '\n' if you want block of commands. 
        """
        
        self.document.add_paragraph(text, style="Configuration")
        
    
    def add_note(self, text):
        """
        Adds a note to the report.
        
        Parameters
        ----------
        text : str
            Message of the note. 
        """
        
        self.document.add_paragraph("Note:", style="Note")
        if text != '':
            note = self.document.add_paragraph(text)
            note.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            note.paragraph_format.left_indent = Pt(48)
            
            
    def add_warning(self, text):
        """
        Adds a warning to the report.
        
        Parameters
        ----------
        text : str
            Message of the warning. 
        """
        
        self.document.add_paragraph("Important:", style="Caution")
        warning = self.document.add_paragraph(text)
        warning.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        warning.paragraph_format.left_indent = Pt(48)
        
        
    def add_bullet(self, text, level=1, numbering=False):
        """
        Adds a bullet paragraph to the report.
        
        Parameters
        ----------
        text : str
            Bullet text.
        level : int
            Supported levels: 1-3.
        numbering : bool
            Show numbers instead of bullets.
        """
        
        if numbering == True:
            bullet = self.document.add_paragraph(text, style="Numbered List")
            bullet.paragraph_format.left_indent = Pt(24*level)
            bullet.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            bullet = self.document.add_paragraph(text, style="List Bullet " + str(level))
            bullet.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            
    def add_table(self, matrix):
        """
        Adds a table to the report.
        
        Parameters
        ----------
        matrix : list
            Maps the matrix list to a table.
        """
        
        cols = len(matrix[0])
        
        table = self.document.add_table(rows=1, cols=cols)
        table.style = 'AS Table'
        
        header_cells = table.rows[0].cells
        for index, cell in enumerate(header_cells):
            cell.text = str(matrix[0][index])
            
        for row in matrix[1:]:
            row_cells = table.add_row().cells
            for index, element in enumerate(row):
                row_cells[index].text = str(element)
                
        self.document.add_paragraph()
    
    
        
    def save(self, filename):
        """
        Exports the structure into a file.
        """
        
        ###### Clean up the unused paragraphs from the template(preface, scope, etc.)
        paragraphs_staged_for_deletion = []
        
        for index, paragraph in enumerate(self.document.paragraphs):
            if '<Enter appropriate text in this section using MSWord style NORMAL.>' in paragraph.text:
                paragraphs_staged_for_deletion.append(index)
                paragraphs_staged_for_deletion.append(index - 1)
                
        for index in sorted(paragraphs_staged_for_deletion, reverse=True):
            delete_element(self.document.paragraphs[index])
        
        self.document.add_page_break()
        
        ###### Append the document ending
        document_ending = Document("template/template.docx")
        
        for paragraph in document_ending.paragraphs[-24:-9]:
            if paragraph.text != '':
                insert_paragraph_into_doc(paragraph, self.document)
        
        self.document.add_page_break()
        self.document.add_paragraph("Document Acceptance", style="Heading 1-No Numbers")
        
        for table in document_ending.tables[-3:]:
            insert_table_into_doc(table, self.document)
            self.document.add_paragraph()
        
        ###### Save the Word file
        self.document.save(filename)
        


def insert_table_into_doc(table, document):
    """
    Inserts an existing Table object into the document.
    """
    inserted_table = document.add_table(rows=len(table.rows), cols=len(table.columns), style=table.style)
    
    for row_idx in range(len(table.rows)):
        for col_idx in range(len(table.columns)):
            inserted_table.cell(row_idx, col_idx).text = table.cell(row_idx, col_idx).text
            
            
def insert_paragraph_into_doc(paragraph, document):
    """
    Inserts an existing Paragraph object into the document.
    """
    inserted_paragraph = document.add_paragraph(paragraph.text, style=paragraph.style)


def delete_element(element):
    """
    Extends the Python DOCX module with the deletion option.
    """
    
    e = element._element
    e.getparent().remove(e)
    e._e = e._element = None
    
